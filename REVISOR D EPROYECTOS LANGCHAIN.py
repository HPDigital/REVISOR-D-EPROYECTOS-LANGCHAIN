"""
REVISOR D EPROYECTOS LANGCHAIN
"""

#!/usr/bin/env python
# coding: utf-8

# In[11]:


from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.schema import SystemMessage, HumanMessage, AIMessage
from langchain.chains import LLMChain


import os
from dotenv import load_dotenv
from docx import Document
import pdfplumber

load_dotenv()

class DocumentAnalyzer:
    def __init__(self, api_key, prompts):
        self.llm = ChatOpenAI(model="gpt-4o",openai_api_key=api_key, temperature=0.3)
        self.prompts = prompts

    def load_document(self, file_path):
        extension = os.path.splitext(file_path)[1].lower()
        if extension == ".pdf":
            return self.load_pdf(file_path)
        elif extension == ".docx":
            return self.load_docx(file_path)
        else:
            raise ValueError(f"Formato de archivo no soportado: {extension}")

    def load_pdf(self, file_path):
        try:
            with pdfplumber.open(file_path) as pdf:
                return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        except Exception as e:
            raise ValueError(f"Error al procesar PDF: {e}")

    def load_docx(self, file_path):
        try:
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
        except Exception as e:
            raise ValueError(f"Error al procesar DOCX: {e}")

    def process_prompt(self, prompt_template, document_text):
        # Crear el prompt
        prompt = PromptTemplate(input_variables=["content"], template=prompt_template)
        # Crear la cadena
        chain = LLMChain(llm=self.llm, prompt=prompt)
        # Ejecutar la cadena
        return chain.run({"content": document_text})

    def analyze_document(self, file_path):
        try:
            document_text = self.load_document(file_path)
            document_text = self.preprocess_text(document_text)
            results = {}
            for title, prompt_template in self.prompts.items():
                print(f"Procesando análisis: {title}")
                results[title] = self.process_prompt(prompt_template, document_text)
            return results
        except Exception as e:
            raise ValueError(f"Error al analizar el documento: {e}")

    def preprocess_text(self, text):
        """Preprocesa el texto para eliminar caracteres innecesarios y normalizarlo."""
        return "\n".join([line.strip() for line in text.splitlines() if line.strip()])

    def save_results(self, output_file, analysis_results):
        try:
            doc = Document()
            doc.add_heading("Resultados del Análisis", level=1)
            for title, result in analysis_results.items():
                doc.add_heading(title, level=2)
                doc.add_paragraph(result)
            doc.save(output_file)
        except Exception as e:
            raise ValueError(f"Error al guardar resultados: {e}")

    def analyze_folder(self, folder_path):
        files = [
            os.path.join(folder_path, f)
            for f in os.listdir(folder_path)
            if os.path.isfile(os.path.join(folder_path, f))
        ]
        for file_path in files:
            try:
                print(f"Procesando archivo: {file_path}")
                results = self.analyze_document(file_path)
                output_file = os.path.splitext(file_path)[0] + "_analysis.docx"
                self.save_results(output_file, results)
                print(f"Análisis completado. Resultados guardados en: {output_file}")
            except Exception as e:
                print(f"Error procesando {file_path}: {e}")

if __name__ == "__main__":
    api_key = os.getenv("OPENAI_API_KEY")
    prompts = {
    "Evaluación de Consistencia": """En tanto que experto en metodología de investigación científica y elaboración de proyectos empresariales, Toma el documento adjunto y extrae la descripción del problema, las causas del problema, los objetivos estratégicos del proyecto y la conclusión. Con base en esta información, analiza la coherencia y consistencia entre: 
        1.	 Extrae el problema central y la conclusión y describe las inconsistencias entre la solución propuesta en la conclusión y el problema central. S La pregunta central es ¿Se esta solucionando el exacto mismo problema en la conclusión que en el problema central?
        2.	 Describe las inconsistencias que hay entre los objetivos estratégicos y causas del problema. 
        3.	Describe las inconsistencias entre la conclusión extraída ofrece una solución efectiva al problema planteado y proporciona retroalimentación detallada sobre cualquier discrepancia lógica. 
        4.	Describe las inconsistencias entre  la propuesta de solución presentada y  la competencia actual, dando solución a las causas mencionadas que originan el problema.""",

    "PESTEL": """En tanto que experto en metodología de investigación científica y elaboración de proyectos empresariales, revisa detalladamente el proyecto adjunto para evaluar si el Análisis PESTEL ha sido correctamente implementado. Específicamente:
        1.	Dimensiones cubiertas: Verifica si todas las dimensiones (Políticas, Económicas, Sociales, Tecnológicas, Ambientales y Legales) están debidamente evaluadas y justificadas con información relevante y actualizada.
        2.	Rigor analítico: Identifica si los análisis en cada dimensión son profundos, bien documentados y basados en fuentes confiables. Evalúa si hay evidencias claras de cómo cada dimensión impacta el entorno externo del proyecto.
        3.	Errores o deficiencias: Señala cualquier falta de análisis, información desactualizada, omisiones importantes o cualquier uso incorrecto de la metodología PESTEL.
        4.	Impacto en el proyecto: Menciona cómo estos posibles errores o deficiencias podrían afectar la calidad general del análisis y la toma de decisiones estratégicas del proyecto.
        Entrega tu análisis en un formato estructurado, señalando concretamente los problemas encontrados (si los hubiera)""",
    "PORTER": """ En tanto que experto en metodología de investigación científica y elaboración de proyectos empresariales, revisa detalladamente el proyecto adjunto para evaluar si el Análisis de las Cinco Fuerzas de Porter ha sido implementado correctamente. Específicamente:
          1.	Cobertura de las cinco fuerzas: Verifica si se han analizado todas las fuerzas correctamente:
              o	Amenaza de nuevos competidores.
              o	Poder de negociación de proveedores.
              o	Poder de negociación de clientes.
              o	Amenaza de productos sustitutos.
              o	Rivalidad entre competidores.
          Comprueba si cada fuerza está descrita con claridad y respaldada por datos o ejemplos relevantes al contexto del proyecto.
          2.	Rigor analítico: Evalúa la profundidad del análisis para cada fuerza. Verifica si:
              o	Se han identificado los factores clave que afectan cada fuerza.
              o	Se ha incluido evidencia actualizada y contextualizada.
              o	Las conclusiones reflejan un entendimiento claro de la dinámica competitiva en el sector analizado.
          3.	Errores o deficiencias: Señala posibles omisiones, análisis superficiales, datos desactualizados, sesgos o malinterpretaciones que afecten la calidad del análisis.
          4.	Impacto en el proyecto: Explica cómo los posibles errores o deficiencias pueden influir en la calidad de las decisiones estratégicas derivadas del análisis de Porter.
          Entrega tu análisis en un formato estructurado, señalando concretamente los problemas encontrados (si los hubiera)""",
    "Investigación de Mercado": """En tanto que experto en metodología de investigación científica y elaboración de proyectos empresariales, revisa detalladamente el proyecto adjunto para evaluar si el Estudio de Mercado está correctamente desarrollado y cumple con los objetivos establecidos. Específicamente:
          1.	Cobertura de los aspectos clave: Verifica si se han analizado con suficiente profundidad los siguientes elementos:
              o	Segmentación de clientes: Examina si los clientes han sido segmentados en base a criterios relevantes (demográficos, psicográficos, geográficos, conductuales) y si se justifica la elección de dichos segmentos.
              o	Tendencias de mercado: Confirma si se identificaron las tendencias actuales en el mercado, con datos relevantes y actualizados.
              o	Competencia: Revisa si se realizó un análisis claro de los principales competidores, incluyendo sus fortalezas, debilidades y estrategias.
              o	Hábitos de consumo: Verifica si se identificaron patrones de consumo relevantes, respaldados por datos confiables y relevantes al contexto del proyecto.
          2.	Rigor analítico: Evalúa si el estudio utiliza fuentes confiables, datos cuantitativos o cualitativos actualizados, y si los análisis reflejan una comprensión profunda del mercado objetivo.
          3.	Errores o deficiencias: Señala omisiones, análisis superficiales, falta de datos relevantes, o errores en la interpretación de información que puedan comprometer la validez del estudio.
          4.	Impacto en el proyecto: Explica cómo los posibles errores o deficiencias en el estudio de mercado podrían afectar la toma de decisiones estratégicas y el éxito del proyecto.
          Entrega tu análisis en un formato estructurado, señalando concretamente los problemas encontrados (si los hubiera)""",
    "Matriz EFE": """En tanto que experto en metodología de investigación científica y elaboración de proyectos empresariales, revisa detalladamente el proyecto adjunto para evaluar si la Matriz EFE (Evaluación de Factores Externos) ha sido correctamente implementada. Específicamente:
          1.	Identificación de factores clave: Verifica si se han identificado las principales oportunidades y amenazas externas que impactan directamente el éxito del proyecto. Evalúa:
              o	Si los factores seleccionados son relevantes, específicos y están respaldados por información objetiva y actualizada.
              o	Si se han considerado fuentes confiables para identificar las oportunidades y amenazas externas.
          2.	Ponderación adecuada: Revisa si los factores clave han sido correctamente ponderados en función de su importancia relativa al entorno externo del proyecto. Comprueba:
              o	Si las ponderaciones suman 1 (o 100%, según el formato utilizado).
              o	Si las asignaciones de peso reflejan correctamente la relevancia e impacto de cada factor en el contexto del proyecto.
          3.	Calificación precisa: Analiza si cada factor ha sido calificado correctamente, considerando la capacidad del proyecto para aprovechar oportunidades o mitigar amenazas. Verifica si las calificaciones están justificadas con ejemplos o análisis claros.
          4.	Cálculo del puntaje ponderado: Comprueba si el puntaje ponderado total se ha calculado correctamente y si la interpretación del resultado es coherente con el análisis realizado.
          5.	Errores o deficiencias: Señala cualquier problema encontrado, como:
              o	Inclusión de factores irrelevantes o falta de factores clave.
              o	Ponderaciones o calificaciones mal justificadas o inconsistentes.
              o	Errores en los cálculos o en la interpretación del puntaje final.
          6.	Impacto en el proyecto: Explica cómo los posibles errores o deficiencias en la Matriz EFE podrían afectar la evaluación estratégica y la toma de decisiones.
          Entrega tu análisis en un formato estructurado, señalando concretamente los problemas encontrados (si los hubiera)""",
    "Matriz EFI": """En tanto que experto en metodología de investigación científica y elaboración de proyectos empresariales, revisa detalladamente el proyecto adjunto para evaluar si la Matriz EFI (Evaluación de Factores Internos) ha sido correctamente desarrollada. Específicamente:
          1.	Identificación de factores internos clave: Verifica si se han identificado adecuadamente las principales fortalezas y debilidades internas de la organización. Evalúa:
              o	Si los factores seleccionados son específicos, relevantes y están basados en evidencia objetiva.
              o	Si los factores reflejan un análisis exhaustivo de los recursos, capacidades, procesos y estructuras internas.
          2.	Ponderación adecuada: Revisa si los factores clave han sido ponderados correctamente en función de su importancia relativa para el éxito del proyecto. Comprueba:
              o	Si las ponderaciones suman 1 (o 100%, según el formato utilizado).
              o	Si las asignaciones de peso reflejan correctamente la relevancia e impacto de cada factor en la capacidad competitiva del proyecto.
          3.	Calificación precisa: Analiza si cada factor ha sido calificado de forma precisa según el desempeño de la organización en relación con ese factor. Verifica si:
              o	Las calificaciones reflejan objetivamente el nivel de fortaleza o debilidad.
              o	Cada calificación está justificada con datos, ejemplos o análisis claros.
          4.	Cálculo del puntaje ponderado: Comprueba si el puntaje ponderado total se ha calculado correctamente y si la interpretación del resultado es consistente con el análisis realizado.
          5.	Errores o deficiencias: Señala cualquier problema encontrado, como:
              o	Inclusión de factores irrelevantes o falta de factores clave.
              o	Ponderaciones o calificaciones mal justificadas o inconsistentes.
              o	Errores en los cálculos o interpretación del puntaje final.
          6.	Impacto en el proyecto: Explica cómo los posibles errores o deficiencias en la Matriz EFI podrían afectar la evaluación estratégica interna y la toma de decisiones.
          Entrega tu análisis en un formato estructurado, señalando concretamente los problemas encontrados (si los hubiera)""",
    "Matriz FODA": """En tanto que experto en metodología de investigación científica y elaboración de proyectos empresariales, revisa detalladamente el proyecto adjunto para evaluar si la Matriz FODA ha sido correctamente implementada y utilizada. Específicamente:
          1.	Identificación de factores internos y externos: Verifica si se han identificado de manera adecuada y específica:
              o	Fortalezas y Debilidades: Analiza si los factores internos reflejan una evaluación profunda de los recursos, capacidades, procesos y limitaciones internas.
              o	Oportunidades y Amenazas: Confirma si los factores externos están basados en tendencias relevantes del entorno, con evidencia clara y actualizada.
          2.	Coherencia en el análisis: Revisa si los factores incluidos son relevantes, están claramente definidos y reflejan la realidad de la organización. Señala si hay omisiones importantes o si algún factor es redundante o irrelevante.
          3.	Cruce estratégico: Evalúa si el cruce entre los factores de la matriz (FO, FA, DO, DA) genera estrategias concretas y accionables. Considera:
              o	Si las estrategias proponen soluciones prácticas para aprovechar oportunidades, mitigar amenazas, potenciar fortalezas y mejorar debilidades.
              o	Si las estrategias están alineadas con los objetivos generales del proyecto.
          4.	Rigor y justificación: Analiza si las estrategias están respaldadas por un razonamiento lógico, con evidencia o datos que sustenten su viabilidad y relevancia.
          5.	Errores o deficiencias: Identifica problemas como:
              o	Factores poco claros, mal definidos o irrelevantes.
              o	Ausencia de estrategias derivadas del análisis cruzado.
              o	Estrategias vagas, poco realistas o no accionables.
          6.	Impacto en el proyecto: Explica cómo los posibles errores o deficiencias en la Matriz FODA podrían afectar la formulación de estrategias y la toma de decisiones del proyecto.
          Entrega tu análisis en un formato estructurado, señalando concretamente los problemas encontrados (si los hubiera)""",
    "Matrices Integradas": """En tanto que experto en metodología de investigación científica y elaboración de proyectos empresariales, revisa detalladamente el proyecto adjunto para evaluar si las Matrices Integradas (Fortalezas-Debilidades / Oportunidades-Amenazas) han sido correctamente implementadas y utilizadas para definir estrategias prioritarias. Específicamente:
          1.	Identificación de factores internos y externos:
              o	Verifica si las fortalezas y debilidades internas reflejan un análisis claro, relevante y basado en datos específicos de la organización.
              o	Comprueba si las oportunidades y amenazas externas están correctamente identificadas y contextualizadas con respecto al entorno en el que opera el proyecto.
          2.	Coherencia en el cruce de factores:
              o	Evalúa si el cruce entre factores internos (fortalezas y debilidades) y externos (oportunidades y amenazas) ha generado combinaciones lógicas y estratégicas.
              o	Revisa si cada cruce (FO, FA, DO, DA) refleja un razonamiento claro y está alineado con los objetivos del proyecto.
          3.	Formulación de estrategias prioritarias:
              o	Analiza si las estrategias resultantes son específicas, accionables y están jerarquizadas según su relevancia e impacto potencial.
              o	Confirma si las estrategias están diseñadas para: 
              	Aprovechar fortalezas para capitalizar oportunidades (FO).
              	Utilizar fortalezas para mitigar amenazas (FA).
              	Superar debilidades para aprovechar oportunidades (DO).
              	Minimizar debilidades frente a amenazas (DA).
          4.	Justificación y viabilidad:
              o	Evalúa si las estrategias propuestas están respaldadas por análisis lógicos y datos suficientes.
              o	Comprueba si las estrategias son realistas y factibles en el contexto del proyecto.
          5.	Errores o deficiencias:
              o	Identifica cualquier problema, como: 
              	Factores internos o externos irrelevantes o mal definidos.
              	Cruces mal estructurados o que no generan estrategias claras.
              	Estrategias genéricas, poco concretas o no priorizadas.
          6.	Impacto en el proyecto:
              o	Explica cómo los posibles errores o deficiencias en las Matrices Integradas podrían afectar la definición de estrategias prioritarias y la toma de decisiones estratégicas del proyecto.
              Entrega tu análisis en un formato estructurado, señalando concretamente los problemas encontrados (si los hubiera)""",
    "Matriz de Perfil Competitivo": """En tanto que experto en metodología de investigación científica y elaboración de proyectos empresariales, revisa detalladamente el proyecto adjunto para evaluar si la Matriz de Perfil Competitivo (MPC) ha sido correctamente implementada y utilizada. Específicamente:
          1.	Identificación de competidores clave:
              o	Verifica si se han identificado los competidores más relevantes y representativos del mercado.
              o	Asegúrate de que la selección de competidores esté justificada y sea relevante en el contexto del proyecto.
          2.	Selección de variables clave:
              o	Evalúa si las variables utilizadas en la MPC (como calidad, precio, sostenibilidad, servicio al cliente, innovación, entre otras) son pertinentes, específicas y críticas para la competitividad en el sector analizado.
              o	Comprueba si cada variable está claramente definida y alineada con los objetivos del proyecto.
          3.	Asignación de pesos:
              o	Revisa si los pesos asignados a cada variable reflejan su importancia relativa en la industria o mercado específico.
              o	Verifica si los pesos suman 1 (o 100%, según el formato utilizado).
          4.	Calificación del desempeño:
              o	Analiza si las calificaciones otorgadas al proyecto y a los competidores son objetivas, coherentes y están respaldadas por datos o análisis relevantes.
              o	Comprueba si las calificaciones reflejan una evaluación precisa del desempeño en cada variable clave.
          5.	Cálculo del puntaje total:
              o	Revisa si los puntajes totales se han calculado correctamente.
              o	Evalúa si la interpretación de los resultados es lógica y consistente con el análisis realizado.
          6.	Comparación y análisis estratégico:
              o	Verifica si la MPC ha sido utilizada para identificar ventajas competitivas, debilidades relativas y áreas de mejora frente a los competidores.
              o	Evalúa si se han derivado conclusiones y recomendaciones estratégicas basadas en los resultados.
          7.	Errores o deficiencias:
              o	Identifica problemas como: 
              	Competidores o variables irrelevantes o mal definidos.
              	Pesos o calificaciones poco claros, mal justificados o inconsistentes.
              	Errores en los cálculos o en la interpretación de los resultados.
          8.	Impacto en el proyecto:
              o	Explica cómo los posibles errores o deficiencias en la MPC podrían afectar la evaluación competitiva y las decisiones estratégicas del proyecto.
                Entrega tu análisis en un formato estructurado, señalando concretamente los problemas encontrados (si los hubiera)""",
    "Análisis Financiero": """En tanto que experto en metodología de investigación científica y elaboración de proyectos empresariales, revisa detalladamente el proyecto adjunto para evaluar si el Análisis Financiero ha sido correctamente implementado y utilizado. Específicamente:
          1.	Evaluación de costos:
              o	Verifica si los costos (fijos y variables) han sido identificados claramente y están respaldados por cálculos precisos y datos relevantes.
              o	Comprueba si los costos están desglosados de manera adecuada y reflejan la realidad del proyecto.
          2.	Ingresos proyectados:
              o	Revisa si los ingresos proyectados están basados en suposiciones realistas y datos confiables.
              o	Evalúa si se han considerado factores como demanda esperada, precios, y ciclos de ingresos en el mercado objetivo.
          3.	Cálculo del punto de equilibrio:
              o	Analiza si el cálculo del punto de equilibrio incluye todos los elementos necesarios, como costos fijos, costos variables y precio de venta.
              o	Verifica si el resultado está correctamente interpretado y es coherente con los objetivos del proyecto.
          4.	Indicadores financieros (TIR y VAN):
              o	Comprueba si la Tasa Interna de Retorno (TIR) y el Valor Actual Neto (VAN) han sido calculados correctamente, con base en flujos de caja proyectados.
              o	Evalúa si los cálculos consideran el costo del capital, la inflación y los periodos adecuados para el análisis.
              o	Verifica si los resultados son interpretados con claridad y se relacionan con la viabilidad del proyecto.
          5.	Coherencia y justificación:
              o	Revisa si todos los datos utilizados (costos, ingresos, tasas de descuento, periodos de análisis, etc.) están debidamente justificados y respaldados por fuentes confiables.
              o	Analiza si las proyecciones y resultados financieros son coherentes con las premisas generales del proyecto.
          6.	Errores o deficiencias:
              o	Identifica problemas como: 
              	Cálculos incorrectos o mal fundamentados.
              	Suposiciones poco realistas o no justificadas.
              	Ausencia de datos clave para evaluar la viabilidad financiera.
          7.	Impacto en el proyecto:
              o	Explica cómo los posibles errores o deficiencias en el Análisis Financiero podrían afectar la toma de decisiones y la evaluación de la viabilidad económica del proyecto.
              Entrega tu análisis en un formato estructurado, señalando concretamente los problemas encontrados (si los hubiera)""",
    "Análisis de Riesgo y Sensibilidad": """En tanto que experto en metodología de investigación científica y elaboración de proyectos empresariales, revisa detalladamente el proyecto adjunto para evaluar si el Análisis de Riesgo y Sensibilidad ha sido correctamente implementado y utilizado. Específicamente:
          1.	Identificación de riesgos clave:
              o	Verifica si se han identificado los riesgos principales que podrían impactar el proyecto.
              o	Evalúa si los riesgos están clasificados y priorizados en función de su probabilidad e impacto.
          2.	Evaluación de escenarios:
              o	Comprueba si se han desarrollado escenarios claros y coherentes: Optimista, Realista y Pesimista.
              o	Verifica si cada escenario está respaldado por datos, suposiciones realistas y un análisis lógico.
          3.	Análisis de sensibilidad:
              o	Evalúa si se han identificado las variables críticas que afectan los resultados del proyecto (por ejemplo, costos, ingresos, tasas de descuento, demanda, etc.).
              o	Revisa si el análisis de sensibilidad muestra cómo los cambios en estas variables impactan los resultados clave (como VAN, TIR, o punto de equilibrio).
          4.	Propuestas de mitigación:
              o	Analiza si se han propuesto estrategias específicas para mitigar los riesgos identificados en cada escenario.
              o	Verifica si las estrategias de mitigación son prácticas, viables y están alineadas con los recursos y capacidades del proyecto.
          5.	Coherencia y justificación:
              o	Revisa si los datos utilizados en el análisis de riesgo y sensibilidad están debidamente justificados y respaldados por fuentes confiables.
              o	Evalúa si las conclusiones del análisis son consistentes con las premisas generales del proyecto.
          6.	Errores o deficiencias:
              o	Identifica problemas como: 
              	Omisión de riesgos clave o escenarios incompletos.
              	Suposiciones poco realistas o no justificadas.
              	Falta de análisis detallado en la sensibilidad de las variables críticas.
              	Ausencia o vaguedad en las estrategias de mitigación propuestas.
          7.	Impacto en el proyecto:
              o	Explica cómo los posibles errores o deficiencias en el análisis de riesgo y sensibilidad podrían comprometer la capacidad del proyecto para anticipar y manejar incertidumbres.
              Entrega tu análisis en un formato estructurado, señalando concretamente los problemas encontrados (si los hubiera)""",

    }
    folder_path = r"C:\Users\HP\Desktop\REVISION DE TRABAJOS  DE DIPLOMADO\TRABAJOS A REVISAR POLEYN HERBAS"

    analyzer = DocumentAnalyzer(api_key, prompts)
    analyzer.analyze_folder(folder_path)


# In[ ]:





# In[ ]:





# In[ ]:





# In[33]:


from langchain.vectorstores import FAISS
from langchain.embeddings.openai import OpenAIEmbeddings
from pydantic import BaseModel  # Importar directamente desde Pydantic
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import faiss
from langchain.docstore.in_memory import InMemoryDocstore

class DocumentAnalyzerWithFAISS:
    def __init__(self, api_key, prompts, faiss_index_path):
        self.api_key = api_key
        self.prompts = prompts
        self.faiss_index_path = faiss_index_path
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        self.embeddings = OpenAIEmbeddings(openai_api_key=api_key)
        self.vector_store = self.load_or_create_faiss_index()

    def load_or_create_faiss_index(self):
        if os.path.exists(self.faiss_index_path):
            # Cargar índice FAISS existente
            return FAISS.load_local(self.faiss_index_path, self.embeddings)
        else:
            # Crear un nuevo índice FAISS
            embedding_example = self.embeddings.embed_query("test")
            if isinstance(embedding_example, list):
                embedding_dim = len(embedding_example)
            else:
                raise TypeError("El embedding devuelto no es una lista válida.")
            index = faiss.IndexFlatL2(embedding_dim)
            docstore = InMemoryDocstore({})  # Cambiado a un Docstore en memoria que soporte adiciones
            index_to_docstore_id = {}
            return FAISS(index=index, docstore=docstore, index_to_docstore_id=index_to_docstore_id, embedding_function=self.embeddings.embed_query)

    def save_faiss_index(self):
        self.vector_store.save_local(self.faiss_index_path)

    def add_documents_to_index(self, documents):
        chunks = []
        for document in documents:
            for chunk in self.text_splitter.split_text(document['content']):
                chunks.append(Document(page_content=chunk, metadata=document['metadata']))
        self.vector_store.add_documents(chunks)

    def load_document(self, file_path):
        extension = os.path.splitext(file_path)[1].lower()
        if extension == ".pdf":
            return self.load_pdf(file_path)
        elif extension == ".docx":
            return self.load_docx(file_path)
        else:
            raise ValueError(f"Formato de archivo no soportado: {extension}")

    def load_pdf(self, file_path):
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        except Exception as e:
            raise ValueError(f"Error al procesar PDF: {e}")

    def load_docx(self, file_path):
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
        except Exception as e:
            raise ValueError(f"Error al procesar DOCX: {e}")

    def analyze_folder(self, folder_path):
        files = [
            os.path.join(folder_path, f)
            for f in os.listdir(folder_path)
            if os.path.isfile(os.path.join(folder_path, f))
        ]
        for file_path in files:
            print(f"Procesando archivo: {file_path}")
            # Carga y agrega documentos al índice
            document_text = self.load_document(file_path)
            metadata = {"file_name": os.path.basename(file_path)}
            self.add_documents_to_index([{"content": document_text, "metadata": metadata}])
        # Guarda el índice FAISS actualizado
        self.save_faiss_index()

if __name__ == "__main__":
    api_key = os.getenv("OPENAI_API_KEY")
    prompts = {
        "Evaluación de Consistencia": """Toma el documento adjunto y extrae la descripción del problema, las causas del problema, los objetivos estratégicos del proyecto y la conclusión. Evalúa si la conclusión aborda el problema identificado y analiza la coherencia lógica entre los objetivos, causas y solución.""",
        "PESTEL": """Revisa si el análisis PESTEL incluye las dimensiones clave (Política, Económica, Social, Tecnológica, Ambiental, Legal) con suficiente profundidad y justificación.""",
    }

    folder_path = r"C:\Users\HP\Desktop\REVISION DE TRABAJOS DE DIPLOMADO\TRABAJOS A REVISAR POLEYN HERBAS"
    faiss_index_path = r"C:\Users\HP\Desktop\FAISS_INDEX"

    analyzer = DocumentAnalyzerWithFAISS(api_key, prompts, faiss_index_path)
    analyzer.analyze_folder(folder_path)


# In[ ]:




